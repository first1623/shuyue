#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成多格式测试文档
"""
from pathlib import Path

# 创建测试目录
test_dir = Path("c:/Users/zhaoy/CodeBuddy/backend/data/test_documents")
test_dir.mkdir(parents=True, exist_ok=True)

# 1. 创建 Markdown 测试文件
md_content = """# 深度学习基础

## 第一章：神经网络概述

深度学习是机器学习的一个分支，它使用多层神经网络来学习数据的表示。

### 1.1 基本概念

神经网络由输入层、隐藏层和输出层组成。每一层包含多个神经元，神经元之间通过权重连接。

### 1.2 激活函数

常用的激活函数包括：
- ReLU
- Sigmoid
- Tanh
- Softmax

## 第二章：卷积神经网络

卷积神经网络（CNN）主要用于图像处理任务。它通过卷积层、池化层和全连接层来提取图像特征。

### 2.1 卷积操作

卷积操作使用卷积核在输入上滑动，提取局部特征。

### 2.2 池化层

池化层用于降低特征图的维度，常见的池化方式有最大池化和平均池化。

## 结论

深度学习在计算机视觉、自然语言处理等领域取得了显著成果。
"""

with open(test_dir / "test.md", "w", encoding="utf-8") as f:
    f.write(md_content)
print(f"[成功] 创建 Markdown 测试文件: {test_dir / 'test.md'}")

# 2. 创建 Word 文档
try:
    from docx import Document
    from docx.shared import Pt
    
    doc = Document()
    doc.add_heading('自然语言处理技术', 0)
    
    doc.add_heading('第一章：NLP 简介', level=1)
    doc.add_paragraph('自然语言处理（NLP）是人工智能的重要分支，研究计算机如何理解和生成人类语言。')
    
    doc.add_heading('1.1 主要任务', level=2)
    doc.add_paragraph('文本分类、情感分析、命名实体识别、机器翻译等。')
    
    doc.add_heading('第二章：深度学习在 NLP 中的应用', level=1)
    doc.add_paragraph('循环神经网络（RNN）和 Transformer 模型在 NLP 任务中取得了突破性进展。')
    
    doc.add_heading('2.1 Word2Vec', level=2)
    doc.add_paragraph('Word2Vec 是一种将词语映射到向量空间的技术，能够捕捉词语之间的语义关系。')
    
    doc.add_heading('2.2 BERT 模型', level=2)
    doc.add_paragraph('BERT 是基于 Transformer 的预训练语言模型，在多项 NLP 任务中取得了最佳性能。')
    
    doc.add_heading('结论', level=1)
    doc.add_paragraph('自然语言处理技术在搜索引擎、智能客服、机器翻译等领域有广泛应用。')
    
    doc.save(test_dir / "test.docx")
    print(f"[成功] 创建 Word 测试文件: {test_dir / 'test.docx'}")
except Exception as e:
    print(f"[失败] 创建 Word 文件失败: {e}")

# 3. 创建 PDF 文件（使用 FPDF）
try:
    from fpdf import FPDF
    
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('SimHei', '', 'C:/Windows/Fonts/simhei.ttf', uni=True)
    pdf.set_font('SimHei', '', 16)
    
    pdf.cell(0, 10, '强化学习基础', ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font('SimHei', '', 12)
    pdf.multi_cell(0, 8, '第一章：强化学习概述')
    pdf.multi_cell(0, 8, '强化学习是一种机器学习方法，智能体通过与环境交互来学习最优策略。')
    pdf.ln(5)
    
    pdf.multi_cell(0, 8, '1.1 核心要素')
    pdf.multi_cell(0, 8, '强化学习包含状态、动作、奖励和策略四个核心要素。')
    pdf.ln(5)
    
    pdf.multi_cell(0, 8, '第二章：Q 学习')
    pdf.multi_cell(0, 8, 'Q 学习是一种无模型的强化学习算法，通过学习状态-动作值函数来找到最优策略。')
    pdf.ln(5)
    
    pdf.multi_cell(0, 8, '结论')
    pdf.multi_cell(0, 8, '强化学习在游戏、机器人控制、自动驾驶等领域有重要应用。')
    
    pdf.output(test_dir / "test.pdf")
    print(f"[成功] 创建 PDF 测试文件: {test_dir / 'test.pdf'}")
except Exception as e:
    print(f"[失败] 创建 PDF 文件失败: {e}")
    # 如果 FPDF 不可用，使用备用方案
    print("提示：PDF 文件需要手动创建或使用其他工具生成")

print("\n测试文档创建完成！")
print(f"测试目录: {test_dir}")
