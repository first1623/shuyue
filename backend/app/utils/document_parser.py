#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档解析器 - 支持多种格式的文档智能解析
集成缓存和重试机制
"""

import os
import io
import logging
from pathlib import Path
from typing import Dict, List, Optional, Union
from datetime import datetime
import tempfile

# PDF解析
import PyPDF2
import pdfplumber

# Word文档解析
from docx import Document
from docx.shared import Inches

# Markdown解析
import markdown
from bs4 import BeautifulSoup

# 文本处理
import chardet
import re

from app.core.config import settings
from app.models.data_overview import DataBookDetail
from app.core.database import get_db

# 导入缓存和重试模块
from app.utils.deepseek_cache import DeepSeekCache, get_cache_instance
from app.utils.deepseek_retry import DeepSeekRetry, RetryConfig, get_retry_instance

logger = logging.getLogger(__name__)

class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.txt': self._parse_txt,
            '.md': self._parse_markdown
        }
        
        # 初始化缓存
        self.cache = DeepSeekCache(
            cache_dir=settings.DEEPSEEK_CACHE_DIR if hasattr(settings, 'DEEPSEEK_CACHE_DIR') else None,
            max_cache_size_mb=settings.DEEPSEEK_CACHE_MAX_SIZE_MB if hasattr(settings, 'DEEPSEEK_CACHE_MAX_SIZE_MB') else 500,
            cache_ttl_hours=settings.DEEPSEEK_CACHE_TTL_HOURS if hasattr(settings, 'DEEPSEEK_CACHE_TTL_HOURS') else 168,
            enable_cache=settings.DEEPSEEK_CACHE_ENABLED if hasattr(settings, 'DEEPSEEK_CACHE_ENABLED') else True
        )
        
        # 初始化重试机制
        retry_config = RetryConfig(
            max_retries=settings.DEEPSEEK_RETRY_MAX_RETRIES if hasattr(settings, 'DEEPSEEK_RETRY_MAX_RETRIES') else 3,
            base_delay=settings.DEEPSEEK_RETRY_BASE_DELAY if hasattr(settings, 'DEEPSEEK_RETRY_BASE_DELAY') else 1.0,
            max_delay=settings.DEEPSEEK_RETRY_MAX_DELAY if hasattr(settings, 'DEEPSEEK_RETRY_MAX_DELAY') else 60.0,
            exponential_base=settings.DEEPSEEK_RETRY_EXPONENTIAL_BASE if hasattr(settings, 'DEEPSEEK_RETRY_EXPONENTIAL_BASE') else 2.0,
            jitter=settings.DEEPSEEK_RETRY_JITTER if hasattr(settings, 'DEEPSEEK_RETRY_JITTER') else True,
            circuit_breaker_enabled=settings.DEEPSEEK_CIRCUIT_BREAKER_ENABLED if hasattr(settings, 'DEEPSEEK_CIRCUIT_BREAKER_ENABLED') else True,
            circuit_failure_threshold=settings.DEEPSEEK_CIRCUIT_FAILURE_THRESHOLD if hasattr(settings, 'DEEPSEEK_CIRCUIT_FAILURE_THRESHOLD') else 5,
            circuit_recovery_timeout=settings.DEEPSEEK_CIRCUIT_RECOVERY_TIMEOUT if hasattr(settings, 'DEEPSEEK_CIRCUIT_RECOVERY_TIMEOUT') else 60,
            request_timeout=settings.DEEPSEEK_REQUEST_TIMEOUT if hasattr(settings, 'DEEPSEEK_REQUEST_TIMEOUT') else 60
        )
        self.retry_handler = DeepSeekRetry(retry_config)
        
        logger.info("DocumentParser 初始化完成，缓存和重试机制已启用")
    
    def parse_document(self, file_path: str, file_id: int) -> Dict:
        """解析文档并提取结构化信息"""
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"不支持的文档格式: {file_extension}")
        
        # 检查文件大小
        file_size = file_path.stat().st_size
        if file_size > settings.MAX_FILE_SIZE:
            raise ValueError(f"文件过大: {file_size} bytes")
        
        try:
            # 调用对应的解析方法
            parser_func = self.supported_formats[file_extension]
            parsed_content = parser_func(file_path)
            
            # 调用DeepSeek API进行智能分析
            ai_analysis = self._call_deepseek_api(parsed_content)
            
            # 保存解析结果到数据库
            result = self._save_parse_result(file_id, parsed_content, ai_analysis)
            
            return result
            
        except Exception as e:
            logger.error(f"文档解析失败 {file_path}: {str(e)}")
            self._save_parse_error(file_id, str(e))
            raise
    
    def _parse_pdf(self, file_path: Path) -> Dict:
        """解析PDF文档"""
        content = {
            "text": "",
            "pages": [],
            "metadata": {},
            "tables": []
        }
        
        try:
            # 使用pdfplumber获得更好的表格提取
            with pdfplumber.open(file_path) as pdf:
                content["metadata"] = {
                    "page_count": len(pdf.pages),
                    "title": pdf.metadata.get('Title', ''),
                    "author": pdf.metadata.get('Author', ''),
                    "subject": pdf.metadata.get('Subject', '')
                }
                
                full_text = []
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    page_tables = page.extract_tables()
                    
                    content["pages"].append({
                        "page_number": i + 1,
                        "text": page_text,
                        "table_count": len(page_tables)
                    })
                    
                    if page_tables:
                        content["tables"].extend(page_tables)
                    
                    full_text.append(page_text)
                
                content["text"] = "\n\n".join(full_text)
                
        except Exception as e:
            logger.warning(f"pdfplumber解析失败，尝试PyPDF2: {e}")
            # 备用方案：使用PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content["metadata"]["page_count"] = len(pdf_reader.pages)
                
                full_text = []
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text() or ""
                    full_text.append(page_text)
                    content["pages"].append({
                        "page_number": i + 1,
                        "text": page_text,
                        "table_count": 0
                    })
                
                content["text"] = "\n\n".join(full_text)
        
        return content
    
    def _parse_docx(self, file_path: Path) -> Dict:
        """解析Word文档"""
        content = {
            "text": "",
            "paragraphs": [],
            "tables": [],
            "metadata": {}
        }
        
        doc = Document(file_path)
        
        # 提取段落
        paragraphs = []
        full_text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():  # 跳过空段落
                paragraphs.append({
                    "text": para.text,
                    "style": para.style.name,
                    "runs": [{"text": run.text, "bold": run.bold, "italic": run.italic} for run in para.runs]
                })
                full_text_parts.append(para.text)
        
        content["paragraphs"] = paragraphs
        content["text"] = "\n".join(full_text_parts)
        
        # 提取表格
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        content["tables"] = tables
        
        # 提取文档属性
        core_props = doc.core_properties
        content["metadata"] = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "word_count": len(content["text"].split())
        }
        
        return content
    
    def _parse_txt(self, file_path: Path) -> Dict:
        """解析纯文本文件"""
        # 检测编码
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
        
        # 读取内容
        with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
            text = file.read()
        
        # 按段落分割（空行分隔）
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        content = {
            "text": text,
            "paragraphs": [{"text": p, "line_count": len(p.split('\n'))} for p in paragraphs],
            "metadata": {
                "encoding": encoding,
                "character_count": len(text),
                "word_count": len(text.split()),
                "line_count": len(text.split('\n')),
                "paragraph_count": len(paragraphs)
            }
        }
        
        return content
    
    def _parse_markdown(self, file_path: Path) -> Dict:
        """解析Markdown文档"""
        # 读取原始内容
        with open(file_path, 'r', encoding='utf-8') as file:
            md_text = file.read()
        
        # 转换为HTML
        html = markdown.markdown(md_text, extensions=['extra', 'codehilite'])
        
        # 解析HTML结构
        soup = BeautifulSoup(html, 'html.parser')
        
        # 提取结构化内容
        sections = []
        current_section = {"level": 0, "title": "", "content": []}
        
        for element in soup.children:
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                
                # 保存前一个section
                if current_section["content"]:
                    sections.append(current_section.copy())
                
                # 开始新section
                current_section = {
                    "level": level,
                    "title": element.get_text().strip(),
                    "content": [],
                    "id": element.get('id', '')
                }
            elif element.name in ['p', 'div', 'blockquote']:
                text = element.get_text().strip()
                if text:
                    current_section["content"].append({
                        "type": element.name,
                        "text": text
                    })
            elif element.name == 'ul':
                items = [li.get_text().strip() for li in element.find_all('li')]
                if items:
                    current_section["content"].append({
                        "type": "list",
                        "items": items
                    })
        
        # 添加最后一个section
        if current_section["content"]:
            sections.append(current_section)
        
        content = {
            "text": md_text,
            "html": html,
            "sections": sections,
            "metadata": {
                "header_count": len([s for s in sections if s["title"]]),
                "paragraph_count": len([c for s in sections for c in s["content"] if c["type"] == "p"]),
                "list_count": len([c for s in sections for c in s["content"] if c["type"] == "list"])
            }
        }
        
        return content
    
    def _call_deepseek_api(self, parsed_content: Dict) -> Dict:
        """调用DeepSeek API进行智能分析（增强版：支持缓存+重试机制+三维度图谱数据提取）"""
        import requests
        import json
        
        text_sample = parsed_content.get("text", "")[:4000]  # 限制长度
        if not text_sample.strip():
            return {
                "abstract": "",
                "keywords": [],
                "theories": [],
                "experiment_flow": "",
                "statistical_methods": [],
                "conclusion": "",
                "confidence_score": 0.0,
                # 三维度图谱新增字段
                "authors": [],
                "theories_used": [],
                "entities": [],
                "entity_relations": []
            }
        
        api_key = settings.DEEPSEEK_API_KEY
        api_base = settings.DEEPSEEK_API_BASE
        model = settings.DEEPSEEK_MODEL
        
        # 1. 检查缓存
        cached_response = self.cache.get(text_sample, model)
        if cached_response:
            logger.info(f"从缓存加载 DeepSeek 响应")
            return cached_response
        
        # 2. API 未配置时使用模拟数据
        if api_key == "your-deepseek-api-key":
            logger.warning("DeepSeek API Key未配置，使用模拟数据")
            return self._mock_deepseek_response(text_sample)
        
        # 3. 准备 API 请求
        url = f"{api_base.rstrip('/')}/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # 增强的提示词，提取三维度图谱数据
        prompt = f"""你是一个专业的学习平台知识图谱系统AI分析助手。请根据下面的文档内容，提取以下信息并以JSON格式返回：
{{
  "abstract": "文档摘要，200字以内",
  "keywords": ["关键词1", "关键词2"],
  "theories": ["涉及的理论或方法名称"],
  "experiment_flow": "实验流程简述",
  "statistical_methods": ["使用的统计或分析方法"],
  "conclusion": "结论简述",
  "authors": ["作者1", "作者2"],
  "theories_used": [
    {{"name": "理论名称", "description": "简短描述"}}
  ],
  "entities": [
    {{"name": "实体名称", "type": "实体类型（如：人名/地名/术语/方法/机构等）", "frequency": 出现次数}}
  ],
  "entity_relations": [
    {{"source": "实体1", "target": "实体2", "relation": "关系类型"}}
  ]
}}

提取要求：
1. authors: 提取文档作者姓名，如有多个作者则全部提取
2. theories_used: 提取文档中明确引用或使用的理论、模型、框架
3. entities: 提取重要实体词，包括：
   - 人名（研究者、学者等）
   - 地名（研究地点等）
   - 术语（专业术语、概念）
   - 方法（研究方法、分析方法）
   - 机构（大学、研究所、公司）
4. entity_relations: 提取实体间的关系，如"使用"、"属于"、"合作"等

文档内容：
{text_sample}"""
        
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": "你是一个精准的信息抽取助手，严格按照要求返回JSON。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.2,
            "max_tokens": 2000
        }
        
        # 4. 使用重试机制调用 API
        def _make_api_request():
            """实际执行 API 请求的内部函数"""
            timeout = settings.DEEPSEEK_REQUEST_TIMEOUT if hasattr(settings, 'DEEPSEEK_REQUEST_TIMEOUT') else 60
            resp = requests.post(url, headers=headers, data=json.dumps(payload), timeout=timeout)
            resp.raise_for_status()
            return resp.json()
        
        try:
            # 使用重试机制执行请求
            data = self.retry_handler.execute_with_retry(_make_api_request)
            content = data["choices"][0]["message"]["content"]
            
            # 清理 Markdown 代码块标记
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0].strip()
            elif "```" in content:
                content = content.split("```")[1].split("```")[0].strip()
            
            # 尝试解析JSON
            try:
                result = json.loads(content)
            except json.JSONDecodeError:
                # 如果返回非JSON，包装成默认结构
                result = {
                    "abstract": content[:200],
                    "keywords": [],
                    "theories": [],
                    "experiment_flow": "",
                    "statistical_methods": [],
                    "conclusion": "",
                    "confidence_score": 0.7,
                    "authors": [],
                    "theories_used": [],
                    "entities": [],
                    "entity_relations": []
                }
            result["confidence_score"] = 0.9
            
            # 5. 保存到缓存
            self.cache.set(text_sample, result, model)
            logger.info("DeepSeek API 响应已缓存")
            
            return result
            
        except Exception as e:
            logger.error(f"DeepSeek API调用失败（已重试）: {e}")
            # 返回模拟数据作为降级处理
            return self._mock_deepseek_response(text_sample)
    
    def _mock_deepseek_response(self, text_sample: str) -> Dict:
        """生成模拟响应（当API不可用时）"""
        return {
            "abstract": self._generate_mock_abstract(text_sample),
            "keywords": self._extract_mock_keywords(text_sample),
            "theories": self._extract_mock_theories(text_sample),
            "experiment_flow": self._extract_mock_experiment_flow(text_sample),
            "statistical_methods": self._extract_mock_statistical_methods(text_sample),
            "conclusion": self._extract_mock_conclusion(text_sample),
            "confidence_score": 0.85,
            # 三维度图谱新增字段
            "authors": self._extract_mock_authors(text_sample),
            "theories_used": self._extract_mock_theories_used(text_sample),
            "entities": self._extract_mock_entities(text_sample),
            "entity_relations": []
        }
    
    def _generate_mock_abstract(self, text: str) -> str:
        """生成模拟摘要（实际应用中应调用真实AI API）"""
        sentences = text.split('。')
        if len(sentences) >= 2:
            return sentences[0] + '。' + sentences[1] + '。'
        return text[:200] + '...' if len(text) > 200 else text
    
    def _extract_mock_keywords(self, text: str) -> List[str]:
        """提取模拟关键词"""
        # 简单的基于词频的关键词提取
        import jieba
        import re
        
        # 清理文本
        clean_text = re.sub(r'[^\w\u4e00-\u9fff]+', ' ', text)
        words = jieba.lcut(clean_text)
        
        # 过滤停用词和短词
        stop_words = {'的', '了', '在', '是', '有', '和', '与', '或', '但', '而', '就', '都', '要', '也', '还', '又', '再', '更', '很', '非常', '可以', '可能', '应该', '需要'}
        filtered_words = [w for w in words if len(w) > 1 and w not in stop_words and not w.isdigit()]
        
        # 统计词频
        from collections import Counter
        word_freq = Counter(filtered_words)
        
        # 返回高频词作为关键词
        keywords = [word for word, freq in word_freq.most_common(10)]
        return keywords[:8]
    
    def _extract_mock_theories(self, text: str) -> List[str]:
        """提取模拟理论引用"""
        # 查找常见的学术理论关键词
        theory_patterns = [
            r'根据\s*([^，。；]+?)\s*理论',
            r'基于\s*([^，。；]+?)\s*方法',
            r'采用\s*([^，。；]+?)\s*模型',
            r'参考\s*([^，。；]+?)\s*研究'
        ]
        
        theories = []
        for pattern in theory_patterns:
            matches = re.findall(pattern, text)
            theories.extend(matches)
        
        return list(set(theories))[:5]  # 去重并限制数量
    
    def _extract_mock_experiment_flow(self, text: str) -> str:
        """提取模拟实验流程"""
        # 查找实验步骤相关的句子
        flow_indicators = ['步骤', '流程', '过程', '实验', '测试', '方法', 'procedure', 'method']
        
        sentences = text.split('。')
        flow_sentences = []
        
        for sentence in sentences:
            if any(indicator in sentence for indicator in flow_indicators):
                flow_sentences.append(sentence.strip())
        
        return '。'.join(flow_sentences[:3]) + '。' if flow_sentences else "未检测到明确的实验流程描述。"
    
    def _extract_mock_statistical_methods(self, text: str) -> List[str]:
        """提取模拟统计方法"""
        stats_methods = [
            '回归分析', '方差分析', '卡方检验', 't检验', '相关性分析',
            '因子分析', '聚类分析', '主成分分析', '判别分析',
            'logistic回归', '线性回归', '多元回归'
        ]
        
        found_methods = [method for method in stats_methods if method in text]
        return found_methods
    
    def _extract_mock_conclusion(self, text: str) -> str:
        """提取模拟结论"""
        # 通常结论出现在文档末尾
        sentences = text.split('。')
        if len(sentences) >= 3:
            # 取最后几个句子作为结论
            conclusion_sentences = sentences[-3:]
            return '。'.join(conclusion_sentences) + '。'
        return text[-300:] + '...' if len(text) > 300 else text
    
    def _extract_mock_authors(self, text: str) -> List[str]:
        """提取模拟作者信息"""
        # 常见作者模式
        author_patterns = [
            r'作者[：:]\s*([^，。\n]+)',
            r'著者[：:]\s*([^，。\n]+)',
            r'撰写[：:]\s*([^，。\n]+)',
            r'([^\s，。]+?)\s*(?:等\s*)?(?:著|编|译)'
        ]
        
        authors = []
        for pattern in author_patterns:
            matches = re.findall(pattern, text[:500])  # 只在前500字查找
            authors.extend(matches)
        
        # 去重并清理
        authors = [a.strip() for a in authors if a.strip() and len(a.strip()) < 20]
        return list(set(authors))[:5]
    
    def _extract_mock_theories_used(self, text: str) -> List[Dict]:
        """提取模拟使用的理论（用于图谱构建）"""
        theories = []
        
        # 常见理论关键词
        theory_keywords = [
            '机器学习', '深度学习', '神经网络', '决策树', '支持向量机',
            '回归分析', '聚类分析', '因子分析', '主成分分析',
            '博弈论', '系统论', '控制论', '信息论',
            '社会网络分析', '内容分析', '扎根理论',
            '现象学', '实证主义', '建构主义'
        ]
        
        for theory in theory_keywords:
            if theory in text:
                theories.append({
                    "name": theory,
                    "description": f"文档中使用了{theory}相关理论或方法"
                })
        
        return theories[:10]
    
    def _extract_mock_entities(self, text: str) -> List[Dict]:
        """提取模拟实体词（用于图谱构建）"""
        import jieba
        from collections import Counter
        
        entities = []
        
        # 清理文本
        clean_text = re.sub(r'[^\w\u4e00-\u9fff]+', ' ', text)
        words = jieba.lcut(clean_text)
        
        # 过滤停用词和短词
        stop_words = {'的', '了', '在', '是', '有', '和', '与', '或', '但', '而', '就', '都', '要', '也', '还', '又', '再', '更', '很', '非常', '可以', '可能', '应该', '需要', '这个', '那个', '这些', '那些', '我们', '他们', '她们', '它们'}
        filtered_words = [w for w in words if len(w) >= 2 and w not in stop_words and not w.isdigit()]
        
        # 统计词频
        word_freq = Counter(filtered_words)
        
        # 实体类型判断
        def get_entity_type(word: str) -> str:
            # 方法关键词
            if any(kw in word for kw in ['分析', '模型', '方法', '算法', '技术']):
                return '方法'
            # 机构关键词
            if any(kw in word for kw in ['大学', '学院', '研究所', '公司', '机构', '中心']):
                return '机构'
            # 概念关键词
            if any(kw in word for kw in ['理论', '概念', '原理', '机制']):
                return '概念'
            return '术语'
        
        # 提取高频词作为实体
        for word, freq in word_freq.most_common(20):
            if freq >= 2:  # 至少出现2次
                entities.append({
                    "name": word,
                    "type": get_entity_type(word),
                    "frequency": freq
                })
        
        return entities[:15]
    
    def _save_parse_result(self, file_id: int, parsed_content: Dict, ai_analysis: Dict) -> Dict:
        """保存解析结果到数据库"""
        try:
            with get_db() as db_session:
                # 检查是否已存在记录
                existing = db_session.query(DataBookDetail).filter(
                    DataBookDetail.file_id == file_id
                ).first()
                
                if existing:
                    # 更新现有记录
                    detail = existing
                else:
                    # 创建新记录
                    detail = DataBookDetail(file_id=file_id)
                    db_session.add(detail)
                
                # 更新基础字段
                detail.abstract = ai_analysis.get("abstract")
                detail.keywords = ai_analysis.get("keywords", [])
                detail.theories = ai_analysis.get("theories", [])
                detail.experiment_flow = ai_analysis.get("experiment_flow")
                detail.statistical_methods = ai_analysis.get("statistical_methods", [])
                detail.conclusion = ai_analysis.get("conclusion")
                detail.full_text = parsed_content.get("text", "")[:10000]  # 限制长度
                detail.parse_status = "completed"
                detail.parse_time = datetime.utcnow()
                detail.deepseek_response = ai_analysis
                
                # 更新三维度图谱字段（新增）
                detail.authors = ai_analysis.get("authors", [])
                detail.theories_used = ai_analysis.get("theories_used", [])
                detail.entities = ai_analysis.get("entities", [])
                detail.entity_relations = ai_analysis.get("entity_relations", [])
                
                db_session.commit()
                
                return {
                    "status": "success",
                    "file_id": file_id,
                    "parse_time": detail.parse_time.isoformat(),
                    "word_count": len(parsed_content.get("text", "").split()),
                    "confidence_score": ai_analysis.get("confidence_score", 0),
                    "deepseek_response": ai_analysis
                }
                
        except Exception as e:
            logger.error(f"保存解析结果失败: {e}")
            # 即使保存失败，也返回分析结果
            return {
                "status": "partial_success",
                "file_id": file_id,
                "error": str(e),
                "word_count": len(parsed_content.get("text", "").split()),
                "confidence_score": ai_analysis.get("confidence_score", 0),
                "deepseek_response": ai_analysis
            }
    
    def _save_parse_error(self, file_id: int, error_message: str):
        """保存解析错误"""
        try:
            with get_db() as db_session:
                existing = db_session.query(DataBookDetail).filter(
                    DataBookDetail.file_id == file_id
                ).first()
                
                if existing:
                    detail = existing
                else:
                    detail = DataBookDetail(file_id=file_id)
                    db_session.add(detail)
                
                detail.parse_status = "failed"
                detail.parse_error = error_message
                detail.parse_time = datetime.utcnow()
                
                db_session.commit()
                
        except Exception as e:
            logger.error(f"保存解析错误失败: {e}")
    
    def get_stats(self) -> Dict:
        """
        获取缓存和重试统计信息
        
        Returns:
            包含缓存和重试统计的字典
        """
        return {
            "cache": self.cache.get_stats(),
            "retry": self.retry_handler.get_stats()
        }
    
    def clear_cache(self):
        """清空缓存"""
        self.cache.clear_all()
        logger.info("已清空 DeepSeek 缓存")
    
    def reset_circuit_breaker(self):
        """重置断路器"""
        self.retry_handler.reset_circuit_breaker()
        logger.info("已重置断路器")